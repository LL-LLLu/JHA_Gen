import streamlit as st
from docx import Document
from openai import OpenAI
from io import BytesIO
import re
import os

# --- CONFIGURATION ---
# IMPORTANT: Make sure this file exists in the same folder as app.py
TEMPLATE_FILENAME = "Template.docx"

# --- PAGE CONFIG ---
st.set_page_config(page_title="JHA Generator", page_icon="🦺", layout="centered")

# --- HELPER FUNCTIONS ---

def clean_response(text):
    """
    Aggressively cleans AI output to ensure consistency.
    Removes 'Output:', 'Response:', quotes, and extra whitespace.
    """
    text = re.sub(r'^(output|response|answer|step \d+|analysis)[:\s-]*', '', text, flags=re.IGNORECASE)
    text = text.replace('"', '').replace("'", "")
    return text.strip()

def get_ai_safety_analysis(client, step_text):
    """
    Sends step to OpenAI with STRICT instructions for consistency.
    """
    try:
        system_msg = "You are a strict data extraction engine for construction safety. You do not chat."
        
        user_msg = (
            f"Analyze this specific MOP step: '{step_text}'\n\n"
            "INSTRUCTIONS:\n"
            "1. DECIDE: Is this step 'Administrative/Safe' OR 'Physical/Hazardous'?\n"
            "   - Safe: Software, checking notes, phone calls, meetings, verifying, notifying.\n"
            "   - Hazardous: Using tools, LOTO, electrical work, ladders, chemicals, pressure.\n"
            "2. OUTPUT FORMAT: Return strictly 'Hazard | Control' (separated by a pipe).\n"
            "3. FOR SAFE STEPS: You MUST return exactly: N/A | N/A\n\n"
            "EXAMPLES:\n"
            "Input: 'Contact the client.' -> Output: N/A | N/A\n"
            "Input: 'Disconnect the main breaker.' -> Output: Electrical Shock | LOTO & Verify Zero Energy\n"
            "Input: 'Update the software tags.' -> Output: N/A | N/A\n"
            "Input: 'Climb ladder to inspect unit.' -> Output: Fall Hazard | Secure Ladder & 3-Points Contact\n\n"
            "Your Output:"
        )
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.0, # Zero creativity = Maximum Consistency
            seed=42          # Mathematical Determinism
        )
        
        raw_content = response.choices[0].message.content.strip()
        cleaned_content = clean_response(raw_content)
        
        if "|" in cleaned_content:
            parts = cleaned_content.split("|")
            return parts[0].strip(), parts[1].strip()
        return "N/A", "N/A"
        
    except Exception as e:
        st.error(f"AI Error: {e}")
        return "Manual Review Required", "Manual Review Required"

def extract_rich_text(cell):
    """Extracts text segments with their original bold/highlight formatting."""
    segments = []
    plain_text_parts = []
    
    for p in cell.paragraphs:
        for run in p.runs:
            text = run.text
            if not text: continue
            plain_text_parts.append(text)
            segments.append({
                "text": text,
                "bold": run.bold,
                "highlight": run.font.highlight_color
            })
    
    return {
        "plain": "".join(plain_text_parts).strip(),
        "segments": segments
    }

def apply_template_font(target_run, template_cell):
    """Clones the base font name and size from the template."""
    if template_cell.paragraphs and template_cell.paragraphs[0].runs:
        tmpl_run = template_cell.paragraphs[0].runs[0]
        target_run.font.name = tmpl_run.font.name
        if tmpl_run.font.size:
            target_run.font.size = tmpl_run.font.size

# --- MAIN APP LOGIC ---

st.title("🦺 MOP to JHA Converter")
st.markdown("Upload your MOP to generate a safety document using the standard template.")

# 1. API KEY HANDLING
api_key = None
if "OPENAI_API_KEY" in st.secrets:
    api_key = st.secrets["OPENAI_API_KEY"]
else:
    with st.sidebar:
        st.header("Settings")
        api_key = st.text_input("OpenAI API Key", type="password")
        st.caption("Key is not stored. Used for this session only.")

# 2. FILE UPLOADER (MOP ONLY)
mop_file = st.file_uploader("Upload MOP (.docx)", type="docx")

# 3. CHECK FOR TEMPLATE
if not os.path.exists(TEMPLATE_FILENAME):
    st.error(f"⚠️ Template file missing! Please ensure '{TEMPLATE_FILENAME}' is in the project folder.")
    st.stop()

# 4. GENERATION PROCESS
if st.button("Generate JHA", type="primary"):
    # Validation
    if not api_key:
        st.error("Please provide an OpenAI API Key.")
        st.stop()
    if not mop_file:
        st.error("Please upload the MOP file.")
        st.stop()

    # Initialize OpenAI
    client = OpenAI(api_key=api_key)

    # Load Documents
    try:
        mop_doc = Document(mop_file)
        # Load Template from local path
        jha_doc = Document(TEMPLATE_FILENAME)
    except Exception as e:
        st.error(f"Error reading files: {e}")
        st.stop()

    # --- STEP A: EXTRACT STEPS FROM MOP ---
    status = st.status("Scanning MOP...", expanded=True)
    steps_data = []
    
    # Deep Scan for Header
    for t_idx, table in enumerate(mop_doc.tables):
        target_col = -1
        header_row = -1
        
        # Search rows 0-6
        for r_idx, row in enumerate(table.rows[:6]):
            for c_idx, cell in enumerate(row.cells):
                clean_txt = cell.text.strip().upper()
                if "DESCRIPTION" in clean_txt and "OPERATION" in clean_txt:
                    target_col = c_idx
                    header_row = r_idx
                    status.write(f"Found Header in Table {t_idx+1}")
                    break
            if target_col != -1: break
            
        if target_col != -1:
            for row in table.rows[header_row + 1:]:
                if len(row.cells) > target_col:
                    cell = row.cells[target_col]
                    txt_check = cell.text.strip()
                    # Filter garbage
                    if (len(txt_check) > 3 
                        and "DESCRIPTION" not in txt_check.upper() 
                        and "DO NOT DELETE" not in txt_check.upper()
                        and "JE DUNN" not in txt_check.upper()):
                        
                        steps_data.append(extract_rich_text(cell))
            break 

    if not steps_data:
        status.update(label="Error: No steps found!", state="error")
        st.error("Could not find 'DESCRIPTION OF WORK OPERATION' column.")
        st.stop()

    status.write(f"Extracted {len(steps_data)} steps.")

    # --- STEP B: PREPARE JHA TEMPLATE ---
    target_table = None
    for table in jha_doc.tables:
        if table.rows and "Sequence" in table.rows[0].cells[0].text:
            target_table = table
            break
            
    if not target_table:
        status.update(label="Error: Template Invalid", state="error")
        st.error("Could not find 'Sequence of Basic Job Steps' table in the Template file.")
        st.stop()

    style_ref_cell = target_table.rows[1].cells[0] if len(target_table.rows) > 1 else None

    # Clear old rows
    for i in range(len(target_table.rows) - 1, 0, -1):
        target_table._element.remove(target_table.rows[i]._element)

    # --- STEP C: AI ANALYSIS LOOP ---
    status.write("Running AI Safety Analysis...")
    progress_bar = st.progress(0)
    
    for i, step_obj in enumerate(steps_data):
        # Update Progress
        progress_bar.progress((i + 1) / len(steps_data))
        
        # 1. AI Analysis
        haz, ctrl = get_ai_safety_analysis(client, step_obj['plain'])
        
        # 2. Add Row
        new_row = target_table.add_row()
        
        # Col 1: Step X + Rich Text
        cell_step = new_row.cells[0]
        cell_step._element.clear_content()
        p = cell_step.add_paragraph()
        
        run_prefix = p.add_run(f"Step {i}:") # Starts at 0
        run_prefix.bold = True
        if style_ref_cell: apply_template_font(run_prefix, style_ref_cell)
        p.add_run("\n")
        
        for seg in step_obj['segments']:
            r = p.add_run(seg['text'])
            if seg['bold']: r.bold = True
            if seg['highlight']: r.font.highlight_color = seg['highlight']
            if style_ref_cell: apply_template_font(r, style_ref_cell)

        # Col 2 & 3: Hazards/Controls
        new_row.cells[1].text = haz
        new_row.cells[2].text = ctrl
        
        if style_ref_cell:
            for col_idx in [1, 2]:
                cell = new_row.cells[col_idx]
                if cell.paragraphs:
                    for run in cell.paragraphs[0].runs:
                        apply_template_font(run, style_ref_cell)

    status.update(label="Complete!", state="complete")
    
    # --- STEP D: DOWNLOAD ---
    buffer = BytesIO()
    jha_doc.save(buffer)
    buffer.seek(0)
    
    st.success("Analysis Complete!")
    st.download_button(
        label="📥 Download Final JHA",
        data=buffer,
        file_name="Final_JHA.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )